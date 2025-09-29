from fastapi import APIRouter, HTTPException, UploadFile, File
import psycopg2
from pydantic import BaseModel
import sys
import os
# 新增資料庫連線管理
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from database_config import get_db_connection, execute_query, execute_transaction
from env_loader import load_env_file

# 載入環境變數
load_env_file()
from typing import Optional, List
from datetime import datetime
import base64
import io
import random
import string
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import fitz  # PyMuPDF
import pandas as pd
import numpy as np
from PIL import Image
import openpyxl
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors

router = APIRouter()

def word_to_pdf_win32com(file_content, filename):
    """使用 win32com 將 Word 轉換為 PDF（最佳格式保留）"""
    try:
        import win32com.client
        import tempfile
        import os
        
        # 建立臨時 Word 檔案
        suffix = '.docx' if filename.lower().endswith('.docx') else '.doc'
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_word:
            temp_word_path = temp_word.name
            temp_word.write(file_content)
            temp_word.flush()
            
        # 建立臨時 PDF 檔案路徑
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
            
        print(f"使用 win32com 轉換 Word: {filename}")
        print(f"臨時 Word: {temp_word_path}")
        print(f"臨時 PDF: {temp_pdf_path}")
        
        try:
            # 創建 Word 應用程序對象
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0  # 不顯示警告
            
            # 打開 Word 文件
            doc = word.Documents.Open(os.path.abspath(temp_word_path))
            
            # 導出為 PDF (17 = wdExportFormatPDF)
            doc.ExportAsFixedFormat(
                OutputFileName=os.path.abspath(temp_pdf_path),
                ExportFormat=17,  # wdExportFormatPDF
                OpenAfterExport=False,
                OptimizeFor=0,  # wdExportOptimizeForSizeOnly
                BitmapMissingFonts=True,
                DocStructureTags=True,
                CreateBookmarks=False,
                UseDocStructureTags=True
            )
            
            # 關閉文件和應用程序
            doc.Close()
            word.Quit()
            
            # 讀取生成的 PDF
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            
            print(f"win32com Word 轉換成功，PDF 大小: {len(pdf_bytes)} bytes")
            
            # 清理臨時檔案
            try:
                os.unlink(temp_word_path)
                os.unlink(temp_pdf_path)
            except:
                pass
            
            return pdf_bytes
            
        except Exception as word_error:
            print(f"Word 應用程式轉換失敗: {word_error}")
            # 確保清理 Word 程序
            try:
                word.Quit()
            except:
                pass
            
            # 清理臨時檔案
            try:
                os.unlink(temp_word_path)
                os.unlink(temp_pdf_path)
            except:
                pass
            
            return None
            
    except ImportError:
        print("win32com 不可用，跳過 Word 轉換")
        return None
    except Exception as e:
        print(f"win32com Word 轉換錯誤: {e}")
        return None

def word_to_pdf_docx(file_content, filename, has_chinese_font):
    """使用 python-docx 讀取 Word 內容並轉換為 PDF"""
    try:
        from docx import Document
        
        # 讀取 Word 文件
        doc = Document(io.BytesIO(file_content))
        
        # 創建 PDF
        buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(buffer, pagesize=A4,
                                  topMargin=0.75*inch, bottomMargin=0.75*inch,
                                  leftMargin=0.75*inch, rightMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # 設定樣式
        if has_chinese_font:
            title_style = ParagraphStyle(
                'DocxTitle',
                parent=styles['Title'],
                fontName='ChineseFont',
                fontSize=16,
                spaceAfter=20,
                alignment=1
            )
            normal_style = ParagraphStyle(
                'DocxNormal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=12,
                leading=16,
                spaceAfter=12
            )
        else:
            title_style = styles['Title']
            normal_style = styles['Normal']
        
        
        # 讀取所有段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 簡單處理粗體文字
                text = paragraph.text
                if paragraph.runs:
                    formatted_text = ""
                    for run in paragraph.runs:
                        if run.bold:
                            formatted_text += f"<b>{run.text}</b>"
                        elif run.italic:
                            formatted_text += f"<i>{run.text}</i>"
                        else:
                            formatted_text += run.text
                    text = formatted_text
                
                para = Paragraph(text, normal_style)
                story.append(para)
        
        # 讀取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            if table_data:
                # 創建 PDF 表格
                pdf_table = Table(table_data)
                if has_chinese_font:
                    table_style = TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'ChineseFont'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ])
                else:
                    table_style = TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ])
                
                pdf_table.setStyle(table_style)
                story.append(pdf_table)
                story.append(Spacer(1, 12))
        
        pdf_doc.build(story)
        return buffer.getvalue()
        
    except ImportError:
        print("python-docx 套件未安裝")
        return None
    except Exception as e:
        print(f"python-docx 轉換失敗: {e}")
        return None

def create_word_placeholder_pdf(filename, has_chinese_font):
    """創建 Word 檔案的佔位符 PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # 使用中文字體樣式
    if has_chinese_font:
        word_style = ParagraphStyle(
            'WordStyle',
            parent=styles['Normal'],
            fontName='ChineseFont',
            fontSize=10,
            leading=14,
        )
        word_title_style = ParagraphStyle(
            'WordTitleStyle',
            parent=styles['Title'],
            fontName='ChineseFont',
            fontSize=16,
            leading=20,
        )
    else:
        word_style = styles['Normal']
        word_title_style = styles['Title']
    
    story = []
    
    content = Paragraph("Word檔案內容已上傳，但無法轉換。請安裝必要套件或檢查檔案格式。", word_style)
    story.append(content)
    
    doc.build(story)
    return buffer.getvalue()

# put
def update_data_to_db(sql_prompt: str, params: tuple = ()):
    """使用新的資料庫連線管理器執行更新操作"""
    try:
        return execute_query(sql_prompt, params, fetch='none')
    except Exception as e:
        print(f"[DB ERROR] {e}")
        raise

# 優化的資料庫更新函數
def update_data_to_db_optimized(sql_prompt: str, params: tuple = ()):
    """優化的資料庫更新函數，使用遠端資料庫連線"""
    try:
        return execute_query(sql_prompt, params, env='remote', fetch='none')
    except Exception as e:
        print(f"[DB ERROR] 遠端資料庫錯誤，回退到本地: {e}")
        # 回退到本地資料庫
        return update_data_to_db(sql_prompt, params)

def query_data_from_db(sql_prompt: str, params: tuple = ()):
    """使用新的資料庫連線管理器執行查詢操作"""
    try:
        return execute_query(sql_prompt, params, fetch='all')
    except Exception as e:
        print(f"[DB ERROR] {e}")
        raise

def check_editor_permission(user_role: str):
    """檢查是否有編輯權限"""
    if user_role != 'editor':
        raise HTTPException(status_code=403, detail="權限不足：僅限編輯者使用此功能")

class RecordUpdate(BaseModel):
    customer_id: Optional[str] = None
    customer_name: Optional[str] = None
    product_id: Optional[str] = None 
    purchase_record: Optional[str] = None
    quantity: Optional[int] = None     
    unit_price: Optional[float] = None    
    amount: Optional[float] = None         
    updated_at: Optional[datetime] = None
    status: Optional[str] = None
    confirmed_by: Optional[str] = None
    confirmed_at: Optional[datetime] = None
    user_role: str

class OrderTransactionCreate(BaseModel):
    customer_id: Optional[str] = None
    product_id: Optional[str] = None
    product_name: Optional[str] = None      
    quantity: Optional[int] = None
    unit_price: Optional[float] = None
    amount: Optional[float] = None
    transaction_date: Optional[str] = None
    user_role: str

# 更新暫存訂單
@router.put("/temp/{id}")
def update_temp(id: int, update_data: RecordUpdate):
    check_editor_permission(update_data.user_role)
    update_fields = update_data.dict(exclude_none=True, exclude={'user_role'})
    
    print(f"更新訂單ID: {id}, 更新欄位: {update_fields}")  # 添加調試
    
    if not update_fields:
        raise HTTPException(status_code=400, detail="沒有提供要更新的欄位")

    set_clause = ", ".join([f"{key} = %s" for key in update_fields])
    sql = f"UPDATE temp_customer_records SET {set_clause} WHERE id = %s"
    params = tuple(update_fields.values()) + (id,)
    
    print(f"執行SQL: {sql}, 參數: {params}")  # 添加調試

    try:
        update_data_to_db(sql, params)
        return {
            "message": "更新成功",
            "id": id,
            "updated_fields": update_fields
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")

# 新增訂單交易記錄
@router.post("/order_transactions")
def create_order_transaction(transaction_data: OrderTransactionCreate):
    check_editor_permission(transaction_data.user_role)
    
    # 生成8個字元的transaction_id
    def generate_transaction_id():
        characters = string.ascii_lowercase + string.digits
        return ''.join(random.choice(characters) for _ in range(8))
    
    transaction_id = generate_transaction_id()
    
    sql = """
    INSERT INTO order_transactions 
    (transaction_id, customer_id, product_id, product_name, quantity, unit_price, amount, transaction_date, currency, document_type) 
    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """
    
    params = (
        transaction_id,
        transaction_data.customer_id,
        transaction_data.product_id,
        transaction_data.product_name,    # 新增這行
        transaction_data.quantity,
        transaction_data.unit_price,
        transaction_data.amount,
        transaction_data.transaction_date,
        'NTD',
        '銷貨'
    )
    
    try:
        update_data_to_db(sql, params)
        return {"message": "交易記錄新增成功"}
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="交易記錄新增失敗")

# 客戶資料管理 - 更新資料
class CustomerUpdate(BaseModel):
    customer_id: Optional[str] = None
    customer_name: Optional[str] = None
    phone_number: Optional[str] = None  # 添加遺漏的欄位
    address: Optional[str] = None
    notes: Optional[str] = None
    delivery_schedule: Optional[str] = None
    is_enabled: Optional[bool] = None
    user_role: str

@router.put("/customer/{customer_id}")
def update_customer(customer_id: str, update_data: CustomerUpdate):
    check_editor_permission(update_data.user_role)
    update_fields = update_data.dict(exclude_none=True, exclude={'user_role'})

    if not update_fields:
        raise HTTPException(status_code=400, detail="沒有提供要更新的欄位")

    # 自動添加 updated_date
    update_fields['updated_date'] = datetime.now().date()

    # 優化：使用更高效的連接管理和預處理語句
    set_clause = ", ".join([f"{key} = %s" for key in update_fields])
    sql = f"UPDATE customer SET {set_clause} WHERE customer_id = %s"
    params = tuple(update_fields.values()) + (customer_id,)

    try:
        # 使用優化的資料庫更新函數
        update_data_to_db_optimized(sql, params)
        
        # 返回更新後的記錄資料（可選）
        return {
            "message": "更新成功",
            "customer_id": customer_id,
            "updated_fields": update_fields,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        print(f"[ERROR] 客戶資料更新失敗: {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")


class ItemSubcategoryUpdate(BaseModel):
    item_id: str
    new_subcategory: str
    user_role: str

@router.put("/product_master/update_subcategory")
def update_item_subcategory(update_data: ItemSubcategoryUpdate):
    check_editor_permission(update_data.user_role)
    
    # 修改 SQL 語句，同時更新 subcategory 和 updated_at
    sql = "UPDATE product_master SET subcategory = %s, updated_at = %s WHERE product_id = %s"
    
    # 使用當前時間作為更新時間
    current_time = datetime.now()
    params = (update_data.new_subcategory, current_time, update_data.item_id)

    try:
        update_data_to_db(sql, params)
        return {
            "message": "品項商品群組更新成功",
            "item_id": update_data.item_id,
            "new_subcategory": update_data.new_subcategory,
            "updated_at": current_time.isoformat()  # 返回更新時間供前端參考
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")
    

# 不活躍客戶處理狀態更新
class InactiveCustomerUpdate(BaseModel):
    processed: Optional[bool] = None
    processed_by: Optional[str] = None
    processed_at: Optional[datetime] = None
    user_role: str

@router.put("/inactive_customer/{customer_name}")
def update_inactive_customer(customer_name: str, update_data: InactiveCustomerUpdate):
    check_editor_permission(update_data.user_role)
    update_fields = update_data.dict(exclude_none=True, exclude={'user_role'})
    
    # 如果沒有提供處理時間，自動設定為當前時間
    if update_fields.get('processed') is True and 'processed_at' not in update_fields:
        update_fields['processed_at'] = datetime.now()
    
    if not update_fields:
        raise HTTPException(status_code=400, detail="沒有提供要更新的欄位")

    set_clause = ", ".join([f"{key} = %s" for key in update_fields])
    sql = f"UPDATE inactive_customers SET {set_clause} WHERE customer_name = %s"
    params = tuple(update_fields.values()) + (customer_name,)

    try:
        update_data_to_db(sql, params)
        return {
            "message": "不活躍客戶狀態更新成功",
            "customer_name": customer_name,
            "updated_fields": update_fields
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")

# 批量更新不活躍客戶處理狀態
class BatchInactiveCustomerUpdate(BaseModel):
    customer_names: List[str]
    processed: bool = True
    processed_by: Optional[str] = None
    user_role: str

@router.put("/inactive_customers/batch_update")
def batch_update_inactive_customers(update_data: BatchInactiveCustomerUpdate):
    check_editor_permission(update_data.user_role)
    if not update_data.customer_names:
        raise HTTPException(status_code=400, detail="沒有提供客戶名稱列表")
    
    try:
        success_count = 0
        failed_customers = []
        
        for customer_name in update_data.customer_names:
            try:
                # 準備更新資料
                processed_at = datetime.now() if update_data.processed else None
                
                sql = """
                UPDATE inactive_customers 
                SET processed = %s, processed_by = %s, processed_at = %s 
                WHERE customer_name = %s
                """
                params = (update_data.processed, update_data.processed_by, processed_at, customer_name)
                
                update_data_to_db(sql, params)
                success_count += 1
                
            except Exception as e:
                print(f"[ERROR] 更新客戶 {customer_name} 失敗: {e}")
                failed_customers.append(customer_name)
        
        return {
            "message": f"批量更新完成",
            "success_count": success_count,
            "total_count": len(update_data.customer_names),
            "failed_customers": failed_customers,
            "updated_fields": {
                "processed": update_data.processed,
                "processed_by": update_data.processed_by,
                "processed_at": datetime.now() if update_data.processed else None
            }
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="批量更新失敗")
    
@router.put("/update_repurchase_reminder/{id}")
def update_reminder_sent(id: int, user_role: str):
    check_editor_permission(user_role)
    sql = "UPDATE repurchase_reminders SET reminder_sent = %s WHERE id = %s"
    params = (True, id)

    try:
        update_data_to_db(sql, params)
        return {
            "message": "提醒狀態更新成功",
            "id": id,
            "reminder_sent": True
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")
    
# 更新repurchase_note
class RepurchaseNoteUpdate(BaseModel):
    repurchase_note: str
    user_role: str

@router.put("/update_repurchase_note/{id}")
def update_repurchase_note(id: int, update_data: RepurchaseNoteUpdate):
    check_editor_permission(update_data.user_role)
    sql = "UPDATE repurchase_reminders SET repurchase_note = %s WHERE id = %s"
    params = (update_data.repurchase_note, id)

    try:
        update_data_to_db(sql, params)
        return {
            "message": "回購備註更新成功",
            "id": id,
            "repurchase_note": update_data.repurchase_note
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")

# RAG 知識庫處理
def setup_chinese_font():
    """設定中文字體支援"""
    try:
        # 嘗試註冊系統中文字體
        # Windows 系統字體路徑
        font_paths = [
            r"C:\Windows\Fonts\msjh.ttc",  # 微軟正黑體
            r"C:\Windows\Fonts\SimHei.ttf",  # 黑體
            r"C:\Windows\Fonts\simsun.ttc",  # 新細明體
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if font_path.endswith('.ttc'):
                        # TTC 字體需要指定子字體索引
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                    else:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    return True
                except:
                    continue
        
        # 如果系統字體都找不到，使用 DejaVu 字體
        return False
    except:
        return False

def convert_file_to_pdf(file_content: bytes, filename: str) -> bytes:
    """將不同格式的檔案轉換為A4直向的PDF"""
    try:
        # 設定中文字體
        has_chinese_font = setup_chinese_font()
        file_extension = filename.lower().split('.')[-1]
        
        # 如果已經是PDF，檢查並調整格式
        if file_extension == 'pdf':
            # 使用PyMuPDF處理PDF
            doc = fitz.open(stream=file_content, filetype="pdf")
            output_doc = fitz.open()
            
            for page in doc:
                # 創建A4大小的新頁面 (595, 842)
                new_page = output_doc.new_page(width=595, height=842)
                
                # 獲取原頁面內容
                page_rect = page.rect
                target_rect = fitz.Rect(0, 0, 595, 842)
                
                # 計算縮放比例，保持長寬比
                scale_x = target_rect.width / page_rect.width
                scale_y = target_rect.height / page_rect.height
                scale = min(scale_x, scale_y)
                
                # 計算居中位置
                scaled_width = page_rect.width * scale
                scaled_height = page_rect.height * scale
                x_offset = (target_rect.width - scaled_width) / 2
                y_offset = (target_rect.height - scaled_height) / 2
                
                # 設置變換矩陣
                mat = fitz.Matrix(scale, scale).pretranslate(x_offset, y_offset)
                new_page.show_pdf_page(target_rect, doc, page.number, mat)
            
            doc.close()
            return output_doc.tobytes()
        
        # 處理Excel檔案
        elif file_extension in ['xls', 'xlsx']:
            try:
                # 根據檔案格式選擇不同的讀取方式
                if file_extension == 'xlsx':
                    # 使用 openpyxl 讀取 .xlsx 檔案
                    workbook = openpyxl.load_workbook(io.BytesIO(file_content))
                    all_sheets_data = {}
                    
                    # 讀取所有工作表
                    for sheet_name in workbook.sheetnames:
                        worksheet = workbook[sheet_name]
                        sheet_data = []
                        max_col = worksheet.max_column
                        max_row = worksheet.max_row
                        
                        # 讀取資料並轉換為字串
                        for row in range(1, max_row + 1):
                            row_data = []
                            for col in range(1, max_col + 1):
                                cell_value = worksheet.cell(row=row, column=col).value
                                if cell_value is None:
                                    cell_value = ""
                                row_data.append(str(cell_value))
                            sheet_data.append(row_data)
                        
                        all_sheets_data[sheet_name] = sheet_data
                        
                else:
                    # 使用 pandas 讀取 .xls 檔案
                    all_sheets = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
                    all_sheets_data = {}
                    
                    for sheet_name, df in all_sheets.items():
                        if not df.empty:
                            # 將 DataFrame 轉換為二維列表，包含標題
                            sheet_data = []
                            # 添加列標題（轉換為字串）
                            column_names = [str(col) for col in df.columns.tolist()]
                            sheet_data.append(column_names)
                            # 添加資料行
                            for _, row in df.iterrows():
                                row_data = []
                                for val in row:
                                    if pd.isna(val):
                                        row_data.append("")
                                    else:
                                        row_data.append(str(val))
                                sheet_data.append(row_data)
                            all_sheets_data[sheet_name] = sheet_data
                        else:
                            all_sheets_data[sheet_name] = []
                
                # 創建PDF
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4,
                                      topMargin=0.5*inch, bottomMargin=0.5*inch,
                                      leftMargin=0.5*inch, rightMargin=0.5*inch)
                story = []
                styles = getSampleStyleSheet()
                
                # 根據是否有中文字體創建樣式
                if has_chinese_font:
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontName='ChineseFont',
                        fontSize=16,
                        spaceAfter=30,
                        alignment=1  # 置中
                    )
                else:
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontSize=16,
                        spaceAfter=30,
                        alignment=1  # 置中
                    )
                
                
                # 處理所有工作表
                for sheet_name, data in all_sheets_data.items():
                    
                    if data:
                        # 建立表格
                        table = Table(data)
                        
                        # 自動調整大小以適應頁面
                        available_width = A4[0] - 2 * 0.5 * inch  # 扣除邊距
                        col_count = len(data[0]) if data else 1
                        col_width = available_width / col_count
                        table._argW = [col_width] * col_count
                        
                        # 設定表格樣式
                        if has_chinese_font:
                            table_style = TableStyle([
                                # 標題行樣式
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'ChineseFont'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                
                                # 資料行樣式
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('FONTNAME', (0, 1), (-1, -1), 'ChineseFont'),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                
                                # 邊框
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                
                                # 交替行背景色
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                            ])
                        else:
                            table_style = TableStyle([
                                # 標題行樣式
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                
                                # 資料行樣式
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                
                                # 邊框
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                
                                # 交替行背景色
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                            ])
                        
                        table.setStyle(table_style)
                        story.append(table)
                    else:
                        # 如果工作表沒有資料
                        if has_chinese_font:
                            empty_style = ParagraphStyle(
                                'EmptyStyle',
                                parent=styles['Normal'],
                                fontName='ChineseFont'
                            )
                        else:
                            empty_style = styles['Normal']
                        
                        empty_para = Paragraph("此工作表無資料", empty_style)
                        story.append(empty_para)
                    
                    # 工作表之間的間距
                    story.append(Spacer(1, 20))
                
                doc.build(story)
                return buffer.getvalue()
                
            except Exception as excel_error:
                print(f"[ERROR] Excel檔案讀取失敗: {excel_error}")
                # 如果Excel讀取失敗，創建一個包含錯誤訊息的PDF
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4)
                styles = getSampleStyleSheet()
                
                # 使用中文字體樣式
                if has_chinese_font:
                    error_style = ParagraphStyle(
                        'ErrorStyle',
                        parent=styles['Normal'],
                        fontName='ChineseFont',
                        fontSize=10,
                        leading=14,
                    )
                    error_title_style = ParagraphStyle(
                        'ErrorTitleStyle',
                        parent=styles['Title'],
                        fontName='ChineseFont',
                        fontSize=16,
                        leading=20,
                    )
                else:
                    error_style = styles['Normal']
                    error_title_style = styles['Title']
                
                story = []
                
                
                error_content = Paragraph(f"檔案讀取失敗：{str(excel_error)}", error_style)
                story.append(error_content)
                story.append(Spacer(1, 0.2*inch))
                
                suggestion = Paragraph("建議：請檢查檔案格式是否正確，或嘗試另存為新的Excel檔案。", error_style)
                story.append(suggestion)
                
                doc.build(story)
                return buffer.getvalue()
        
        # 處理Word檔案
        elif file_extension in ['doc', 'docx']:
            try:
                # 先嘗試使用 win32com 保持原始版面
                pdf_content = word_to_pdf_win32com(file_content, filename)
                if pdf_content:
                    return pdf_content

                # win32com 不可用時，對 .docx 使用 python-docx
                if file_extension == 'docx':
                    pdf_content = word_to_pdf_docx(file_content, filename, has_chinese_font)
                    if pdf_content:
                        return pdf_content

                # 無法轉換時提供佔位內容
                return create_word_placeholder_pdf(filename, has_chinese_font)

            except Exception as word_error:
                print(f"[ERROR] Word 檔案處理失敗: {word_error}")
                return create_word_placeholder_pdf(filename, has_chinese_font)

        else:
            raise HTTPException(status_code=400, detail=f"不支援的檔案格式: {file_extension}。支援的格式: pdf, doc, docx, xls, xlsx")
            
    except Exception as e:
        print(f"[ERROR] 檔案轉換失敗: {e}")
        error_detail = f"資料轉換失敗：{str(e)}"
        if "Excel" in str(e) or "openpyxl" in str(e) or "xlrd" in str(e):
            error_detail = f"Excel檔案讀取失敗：{str(e)}。請檢查檔案是否損壞或格式是否正確。"
        raise HTTPException(status_code=500, detail=error_detail)

class RAGKnowledgeBase(BaseModel):
    title: str
    text_content: Optional[str] = None
    files: Optional[List[dict]] = None  # 包含檔名和base64內容
    delete_file_index: Optional[int] = None  # 要刪除的檔案索引
    user_role: str

@router.put("/rag/save_knowledge")
def save_rag_knowledge(knowledge_data: RAGKnowledgeBase):
    """儲存RAG知識庫內容，包含文字和檔案"""
    check_editor_permission(knowledge_data.user_role)
    try:
        # 準備PDF檔案數據
        pdf_files = []
        
        if knowledge_data.files:
            for file_info in knowledge_data.files:
                filename = file_info.get('filename')
                base64_content = file_info.get('content', '')
                frontend_converted = file_info.get('frontend_converted', False)
                
                # 解碼檔案內容 - 檢查是否有 data URL 前綴
                if base64_content.startswith('data:'):
                    # 移除 data URL 前綴
                    base64_data = base64_content.split(',', 1)[1]
                else:
                    base64_data = base64_content
                
                file_content = base64.b64decode(base64_data)
                
                # 如果是前端已轉換的檔案，直接使用內容
                if frontend_converted:
                    pdf_content = file_content
                else:
                    # 非前端轉換的檔案，使用後端轉換
                    pdf_content = convert_file_to_pdf(file_content, filename)
                
                # 計算內容雜湊用於除錯
                import hashlib
                content_hash = hashlib.md5(pdf_content).hexdigest()[:8]
                
                pdf_files.append({
                    'filename': filename,
                    'pdf_content': pdf_content
                })
        
        # 檢查是否已存在該標題的記錄
        check_sql = "SELECT COUNT(*) FROM rag WHERE title = %s"
        result = query_data_from_db(check_sql, (knowledge_data.title,))
        exists = result[0][0] > 0
        
        if exists:
            # 更新現有記錄
            if knowledge_data.delete_file_index is not None:
                # 刪除指定索引的檔案
                get_existing_sql = "SELECT file_content, file_name FROM rag WHERE title = %s"
                existing_data = query_data_from_db(get_existing_sql, (knowledge_data.title,))
                
                if existing_data and existing_data[0]:
                    existing_file_content = list(existing_data[0][0]) if existing_data[0][0] else []
                    existing_file_names = list(existing_data[0][1]) if existing_data[0][1] else []
                    
                    # 刪除指定索引的檔案
                    if 0 <= knowledge_data.delete_file_index < len(existing_file_names):
                        del existing_file_content[knowledge_data.delete_file_index]
                        del existing_file_names[knowledge_data.delete_file_index]
                        
                        sql = """
                        UPDATE rag 
                        SET text_content = %s, file_content = %s, file_name = %s
                        WHERE title = %s
                        """
                        params = (
                            [knowledge_data.text_content] if knowledge_data.text_content else [],
                            existing_file_content,
                            existing_file_names,
                            knowledge_data.title
                        )
                        update_data_to_db(sql, params)
            elif pdf_files:
                # 先獲取現有的檔案內容
                get_existing_sql = "SELECT file_content, file_name FROM rag WHERE title = %s"
                existing_data = query_data_from_db(get_existing_sql, (knowledge_data.title,))
                
                existing_file_content = existing_data[0][0] if existing_data and existing_data[0][0] else []
                existing_file_names = existing_data[0][1] if existing_data and existing_data[0][1] else []
                
                # 累積新檔案到現有檔案中
                new_file_content = list(existing_file_content) if existing_file_content else []
                new_file_names = list(existing_file_names) if existing_file_names else []
                
                for file_info in pdf_files:
                    # 檢查是否已經存在相同檔名，如果存在就跳過
                    if file_info['filename'] not in new_file_names:
                        hex_content = file_info['pdf_content'].hex()
                        new_file_content.append(hex_content)
                        new_file_names.append(file_info['filename'])
                
                sql = """
                UPDATE rag 
                SET text_content = %s, file_content = %s, file_name = %s
                WHERE title = %s
                """
                params = (
                    [knowledge_data.text_content] if knowledge_data.text_content else [],
                    new_file_content,
                    new_file_names,
                    knowledge_data.title
                )
                update_data_to_db(sql, params)
            elif knowledge_data.files is None:
                # 如果files明確為None，清除檔案內容
                sql = """
                UPDATE rag 
                SET text_content = %s, file_content = NULL, file_name = NULL
                WHERE title = %s
                """
                params = ([knowledge_data.text_content] if knowledge_data.text_content else [], knowledge_data.title)
                update_data_to_db(sql, params)
            else:
                # 只更新文字內容，不動檔案
                sql = """
                UPDATE rag 
                SET text_content = %s
                WHERE title = %s
                """
                params = ([knowledge_data.text_content] if knowledge_data.text_content else [], knowledge_data.title)
                update_data_to_db(sql, params)
        else:
            # 新增記錄
            if pdf_files:
                # 儲存所有檔案到陣列中
                file_contents = []
                file_names = []
                
                for file_info in pdf_files:
                    file_contents.append(file_info['pdf_content'].hex())
                    file_names.append(file_info['filename'])
                
                sql = """
                INSERT INTO rag (title, text_content, file_content, file_name)
                VALUES (%s, %s, %s, %s)
                """
                params = (
                    knowledge_data.title,
                    [knowledge_data.text_content] if knowledge_data.text_content else [],
                    file_contents,
                    file_names
                )
                update_data_to_db(sql, params)
            else:
                # 只有文字內容的記錄
                sql = """
                INSERT INTO rag (title, text_content)
                VALUES (%s, %s)
                """
                params = (
                    knowledge_data.title,
                    [knowledge_data.text_content] if knowledge_data.text_content else []
                )
                update_data_to_db(sql, params)
        
        return {
            "message": "知識庫儲存成功",
            "title": knowledge_data.title,
            "files_processed": len(pdf_files) if pdf_files else 0
        }
        
    except Exception as e:
        print(f"[ERROR] RAG儲存失敗: {e}")
        raise HTTPException(status_code=500, detail=f"知識庫儲存失敗: {str(e)}")

@router.put("/rag/delete_knowledge/{title}")
def delete_rag_knowledge(title: str, user_role: str):
    """刪除RAG知識庫條目"""
    check_editor_permission(user_role)
    try:
        sql = "DELETE FROM rag WHERE title = %s"
        params = (title,)
        
        update_data_to_db(sql, params)
        
        return {
            "message": "知識庫條目刪除成功",
            "title": title
        }
        
    except Exception as e:
        print(f"[ERROR] RAG刪除失敗: {e}")
        raise HTTPException(status_code=500, detail="知識庫條目刪除失敗")

class RAGTitleUpdate(BaseModel):
    old_title: str
    new_title: str
    user_role: str

@router.put("/rag/update_title")
def update_rag_title(update_data: RAGTitleUpdate):
    """更新RAG知識庫條目標題"""
    check_editor_permission(update_data.user_role)
    try:
        # 檢查新標題是否已存在
        check_sql = "SELECT COUNT(*) FROM rag WHERE title = %s AND title != %s"
        result = query_data_from_db(check_sql, (update_data.new_title, update_data.old_title))
        
        if result[0][0] > 0:
            raise HTTPException(status_code=400, detail="新標題已存在，請使用其他標題")
        
        # 更新標題
        sql = "UPDATE rag SET title = %s WHERE title = %s"
        params = (update_data.new_title, update_data.old_title)
        
        update_data_to_db(sql, params)
        
        return {
            "message": "標題更新成功",
            "old_title": update_data.old_title,
            "new_title": update_data.new_title
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] RAG標題更新失敗: {e}")
        raise HTTPException(status_code=500, detail="標題更新失敗")
    
class RestockPredictionStatusUpdate(BaseModel):
    prediction_id: int
    prediction_status: str
    user_role: str

# 更新補貨預測狀態
@router.put("/update_restock_prediction_status")
def update_restock_prediction_status(update_data: RestockPredictionStatusUpdate):
    check_editor_permission(update_data.user_role)
    
    # 驗證 prediction_status 值
    if update_data.prediction_status not in ['fulfilled', 'cancelled']:
        raise HTTPException(status_code=400, detail="prediction_status 必須是 'fulfilled' 或 'cancelled'")
    
    try:
        # 根據 prediction_id 更新 prophet_predictions 表
        sql = """
        UPDATE prophet_predictions 
        SET prediction_status = %s 
        WHERE prediction_id = %s
        """
        params = (update_data.prediction_status, update_data.prediction_id)
        
        update_data_to_db(sql, params)
        
        return {
            "message": "補貨預測狀態更新成功",
            "prediction_id": update_data.prediction_id,
            "prediction_status": update_data.prediction_status
        }
        
    except Exception as e:
        print(f"[ERROR] 補貨預測狀態更新失敗: {e}")
        raise HTTPException(status_code=500, detail="補貨預測狀態更新失敗")

class DeliveryScheduleUpdate(BaseModel):
    customer_id: Optional[str] = None
    customer_name: Optional[str] = None
    product_name: Optional[str] = None
    status: Optional[str] = None
    quantity: Optional[int] = None
    delivery_date: Optional[datetime] = None
    updated_at: Optional[datetime] = None

# 配送計劃管理 - 更新資料
@router.put("/delivery_schedule/{record_id}")
def update_delivery_schedule(record_id: int, update_data: DeliveryScheduleUpdate):
    update_fields = update_data.dict(exclude_none=True)
    
    if not update_fields:
        raise HTTPException(status_code=400, detail="沒有提供要更新的欄位")
    
    # 自動更新 updated_at 欄位
    update_fields['updated_at'] = datetime.now()
    
    set_clause = ", ".join([f"{key} = %s" for key in update_fields])
    sql = f"UPDATE delivery_schedule SET {set_clause} WHERE id = %s"
    params = tuple(update_fields.values()) + (record_id,)
    
    try:
        update_data_to_db(sql, params)
        return {
            "message": "配送計劃更新成功",
            "id": record_id,
            "updated_fields": update_fields
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="資料庫更新失敗")

# 批量更新配送狀態
@router.put("/delivery_schedule/batch_update_status")
def batch_update_delivery_status(record_ids: List[int], new_status: str):
    try:
        placeholders = ",".join(["%s"] * len(record_ids))
        sql = f"""
        UPDATE delivery_schedule 
        SET status = %s, updated_at = %s 
        WHERE id IN ({placeholders})
        """
        params = [new_status, datetime.now()] + record_ids
        
        update_data_to_db(sql, tuple(params))
        return {
            "message": f"批量更新 {len(record_ids)} 筆配送狀態成功",
            "updated_count": len(record_ids),
            "new_status": new_status
        }
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="批量更新失敗")
    
# 新增客戶創建的 Pydantic 模型和 API 端點
class CustomerCreate(BaseModel):
    customer_id: str
    customer_name: str
    phone_number: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    district: Optional[str] = None
    notes: Optional[str] = None
    delivery_schedule: Optional[str] = None
    line_id: Optional[str] = None
    is_enabled: Optional[bool] = True 
    user_role: str

@router.post("/create_customer")
def create_customer(customer_data: CustomerCreate):
    check_editor_permission(customer_data.user_role)
    
    try:
        # 準備事務查詢
        queries_params = []

        # 創建客戶 - 新增 is_enabled 和 updated_date 欄位
        customer_sql = """
        INSERT INTO customer
        (customer_id, customer_name, phone_number, address, city, district, notes, delivery_schedule, line_id, is_enabled, updated_date)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        customer_params = (
            customer_data.customer_id,
            customer_data.customer_name,
            customer_data.phone_number,
            customer_data.address,
            customer_data.city,
            customer_data.district,
            customer_data.notes,
            customer_data.delivery_schedule,
            customer_data.line_id,
            customer_data.is_enabled if customer_data.is_enabled is not None else True,  # 預設為 True
            datetime.now().date()  # 設定為今天日期
        )
        queries_params.append((customer_sql, customer_params))

        # 如果有 line_id，則寫入 customer_line_mapping 表
        if customer_data.line_id:
            mapping_sql = """
            INSERT INTO customer_line_mapping (customer_id, line_id, created_date, notes)
            VALUES (%s, %s, %s, %s)
            ON CONFLICT (customer_id, line_id) DO NOTHING
            """
            mapping_params = (
                customer_data.customer_id,
                customer_data.line_id,
                datetime.now(),
                f"從新訂單系統創建客戶時建立對應關係 - {customer_data.customer_name}"
            )
            queries_params.append((mapping_sql, mapping_params))

        # 執行事務
        execute_transaction(queries_params)
        
        return {"message": "客戶創建成功", "customer_id": customer_data.customer_id}
    except Exception as e:
        print(f"[ERROR] 客戶創建失敗: {e}")
        raise HTTPException(status_code=500, detail="客戶創建失敗")

class TempOrderCreate(BaseModel):
    customer_id: Optional[str] = None
    customer_name: Optional[str] = None
    line_id: Optional[str] = None
    product_id: Optional[str] = None
    purchase_record: str
    quantity: Optional[int] = None
    unit_price: Optional[float] = None
    amount: Optional[float] = None
    conversation_record: str = "手動新增訂單"
    order_note: Optional[str] = None
    label: str = "ORDER"
    is_new_product: bool = False
    status: str = "0"  # 預設為未確認
    confirmed_by: Optional[str] = None  
    confirmed_at: Optional[str] = None
    user_role: str

@router.post("/create_temp_order")
def create_temp_order(order_data: TempOrderCreate):
    check_editor_permission(order_data.user_role)
    
    current_time = datetime.now().isoformat()
    
    try:
        # 只插入到 temp_customer_records
        temp_sql = """
        INSERT INTO temp_customer_records
        (customer_id, customer_name, line_id, product_id, purchase_record, quantity,
        unit_price, amount, conversation_record, order_note, label, is_new_product,
        status, confirmed_by, confirmed_at, created_at, updated_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """

        normalized_is_new_product = bool(order_data.is_new_product)
        if isinstance(order_data.is_new_product, str):
            normalized_is_new_product = order_data.is_new_product.lower() == 'true'

        if order_data.status == "1":
            if not order_data.product_id:
                raise HTTPException(status_code=400, detail="缺少 product_id，無法確認是否為新品項")

            check_existing_sql = """
                SELECT 1
                FROM order_transactions
                WHERE customer_id = %s AND product_id = %s AND status = 1
                LIMIT 1
            """
            existing_purchase = execute_query(
                check_existing_sql,
                (order_data.customer_id, order_data.product_id),
                fetch='one'
            )

            if existing_purchase:
                normalized_is_new_product = False
            else:
                normalized_is_new_product = True

        order_data.is_new_product = normalized_is_new_product


        temp_params = (
            order_data.customer_id,
            order_data.customer_name,
            order_data.line_id,
            order_data.product_id,
            order_data.purchase_record,
            order_data.quantity,
            order_data.unit_price,
            order_data.amount,
            order_data.conversation_record,
            order_data.order_note,
            order_data.label,
            order_data.is_new_product,
            order_data.status,
            order_data.confirmed_by,
            order_data.confirmed_at,
            current_time,
            current_time
        )

        execute_query(temp_sql, temp_params, fetch='none')
        
        return {"message": "訂單新增成功"}
        
    except Exception as e:
        print(f"[ERROR] {e}")
        raise HTTPException(status_code=500, detail="訂單新增失敗")
    
# 新增客戶Line對應關係的 Pydantic 模型
class CustomerLineMappingCreate(BaseModel):
    customer_id: str
    line_id: str
    user_role: str

@router.post("/customer_line_mapping")
def create_customer_line_mapping(mapping_data: CustomerLineMappingCreate):
    check_editor_permission(mapping_data.user_role)
    
    try:
        current_time = datetime.now().isoformat()
        
        # 檢查是否已存在對應關係
        check_sql = "SELECT COUNT(*) FROM customer_line_mapping WHERE customer_id = %s AND line_id = %s"
        result = execute_query(check_sql, (mapping_data.customer_id, mapping_data.line_id), fetch='one')
        exists = result[0] > 0

        if not exists:
            # 新增對應關係
            insert_sql = """
            INSERT INTO customer_line_mapping (customer_id, line_id, created_date, notes)
            VALUES (%s, %s, %s, %s)
            """
            execute_query(insert_sql, (
                mapping_data.customer_id,
                mapping_data.line_id,
                current_time,
                f"由 {mapping_data.user_role} 創建"
            ), fetch='none')

            return {"success": True, "message": "客戶Line對應關係已建立"}
        else:
            return {"success": True, "message": "對應關係已存在"}
                    
    except Exception as e:
        print(f"[ERROR] create_customer_line_mapping: {e}")
        raise HTTPException(status_code=500, detail="建立對應關係失敗")

# 滯銷品狀態更新
class SalesChangeStatusUpdate(BaseModel):
    product_name: str
    status: bool = True
    processed_by: Optional[str] = None
    user_role: str

@router.put("/update_sales_change_status")
def update_sales_change_status(update_data: SalesChangeStatusUpdate):
    check_editor_permission(update_data.user_role)

    try:
        # 根據商品名稱更新 sales_change_table 的狀態
        sql = """
        UPDATE sales_change_table
        SET status = %s, processed_by = %s, processed_at = %s
        WHERE product_id IN (
            SELECT product_id FROM product_master WHERE name_zh = %s
        )
        """
        params = (
            update_data.status,
            update_data.processed_by,
            datetime.now(),
            update_data.product_name
        )

        update_data_to_db(sql, params)

        return {
            "message": "滯銷品狀態更新成功",
            "product_name": update_data.product_name,
            "status": update_data.status,
            "processed_by": update_data.processed_by
        }

    except Exception as e:
        print(f"[ERROR] 滯銷品狀態更新失敗: {e}")
        raise HTTPException(status_code=500, detail="滯銷品狀態更新失敗")

# 滯銷品狀態更新 - 使用 product_id
class SalesChangeStatusUpdateById(BaseModel):
    product_id: str
    status: bool = True
    processed_by: Optional[str] = None
    user_role: str

@router.put("/update_sales_change_status_by_id")
def update_sales_change_status_by_id(update_data: SalesChangeStatusUpdateById):
    check_editor_permission(update_data.user_role)

    try:
        # 只更新 status 欄位，不更新 processed_by 和 processed_at
        sql = """
        UPDATE sales_change_table
        SET status = %s
        WHERE product_id = %s
        """
        params = (
            update_data.status,
            update_data.product_id
        )

        update_data_to_db(sql, params)

        return {
            "message": "滯銷品狀態更新成功",
            "product_id": update_data.product_id,
            "status": update_data.status
        }

    except Exception as e:
        print(f"[ERROR] 滯銷品狀態更新失敗: {e}")
        raise HTTPException(status_code=500, detail="滯銷品狀態更新失敗")

# === LINE Bot 設定管理 ===
import re

def read_line_env_variable(key: str) -> str:
    """讀取 LINE Bot .env 檔案中的特定變數"""
    env_file_path = "/home/chou_fish_988/Documents/988/Line_bot/.env"
    try:
        with open(env_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    env_key, env_value = line.split('=', 1)
                    if env_key.strip() == key:
                        return env_value.strip()
        return ""
    except Exception as e:
        print(f"[ERROR] 讀取 .env 檔案失敗: {e}")
        raise HTTPException(status_code=500, detail="讀取設定檔失敗")

def update_line_env_variable(key: str, value: str) -> bool:
    """更新 LINE Bot .env 檔案中的特定變數"""
    env_file_path = "/home/chou_fish_988/Documents/988/Line_bot/.env"
    try:
        # 讀取現有內容
        with open(env_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # 尋找並更新指定的變數
        updated = False
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if line_stripped and not line_stripped.startswith('#') and '=' in line_stripped:
                env_key = line_stripped.split('=', 1)[0].strip()
                if env_key == key:
                    lines[i] = f"{key}={value}\n"
                    updated = True
                    break

        # 如果變數不存在，添加到檔案末尾
        if not updated:
            lines.append(f"{key}={value}\n")

        # 寫回檔案
        with open(env_file_path, 'w', encoding='utf-8') as f:
            f.writelines(lines)

        return True
    except Exception as e:
        print(f"[ERROR] 更新 .env 檔案失敗: {e}")
        raise HTTPException(status_code=500, detail="更新設定檔失敗")

def reload_line_env():
    """通知 LINE Bot 重新載入環境變數"""
    try:
        # 這裡可以發送信號給 LINE Bot 重新載入 .env
        # 目前簡單地更新系統環境變數
        env_file_path = "/home/chou_fish_988/Documents/988/Line_bot/.env"
        with open(env_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    os.environ[key.strip()] = value.strip()
        return True
    except Exception as e:
        print(f"[ERROR] 重新載入環境變數失敗: {e}")
        return False

class LineReplyToggleRequest(BaseModel):
    enabled: bool
    user_role: str

@router.get("/line-reply-status")
def get_line_reply_status():
    """獲取 LINE Bot 回覆開關狀態"""
    try:
        current_value = read_line_env_variable("ENABLE_LINE_REPLY")
        enabled = current_value.lower() == 'true'
        return {"enabled": enabled}
    except Exception as e:
        print(f"[ERROR] 獲取 LINE 回覆狀態失敗: {e}")
        raise HTTPException(status_code=500, detail="獲取設定失敗")

@router.put("/line-reply-status")
def update_line_reply_status(request: LineReplyToggleRequest):
    """更新 LINE Bot 回覆開關狀態"""
    check_editor_permission(request.user_role)
    try:
        # 更新 .env 檔案
        value = "true" if request.enabled else "false"
        success = update_line_env_variable("ENABLE_LINE_REPLY", value)

        if success:
            # 重新載入環境變數
            reload_line_env()
            return {
                "message": "LINE Bot 回覆設定更新成功",
                "enabled": request.enabled
            }
        else:
            raise HTTPException(status_code=500, detail="設定更新失敗")

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] 更新 LINE 回覆狀態失敗: {e}")
        raise HTTPException(status_code=500, detail="設定更新失敗")

# 新增 Pydantic 模型
class CustomerDeleteRequest(BaseModel):
    user_role: str

# 修改刪除端點
@router.delete("/customer/{customer_id}")
def delete_customer(customer_id: str, request: CustomerDeleteRequest):
    check_editor_permission(request.user_role)
    
    try:
        # 檢查客戶是否存在
        check_sql = "SELECT COUNT(*) FROM customer WHERE customer_id = %s"
        result = query_data_from_db(check_sql, (customer_id,))
        if not result or result[0][0] == 0:
            raise HTTPException(status_code=404, detail="客戶不存在")
        
        # 刪除客戶
        delete_sql = "DELETE FROM customer WHERE customer_id = %s"
        update_data_to_db(delete_sql, (customer_id,))
        
        return {
            "message": "客戶刪除成功",
            "customer_id": customer_id
        }
    except Exception as e:
        print(f"[ERROR] 客戶刪除失敗: {e}")
        raise HTTPException(status_code=500, detail="客戶刪除失敗")

# 產品刪除請求模型
class ProductDeleteRequest(BaseModel):
    product_id: str
    user_role: str

# 產品刪除 API
@router.delete("/product/delete")
def delete_product(request: ProductDeleteRequest):
    check_editor_permission(request.user_role)

    try:
        product_id = request.product_id

        # 檢查產品是否存在於 product_master
        check_sql = "SELECT COUNT(*) FROM product_master WHERE product_id = %s"
        result = query_data_from_db(check_sql, (product_id,))
        if not result or result[0][0] == 0:
            raise HTTPException(status_code=404, detail="產品不存在")

        # 從 inventory 表中刪除
        delete_inventory_sql = "DELETE FROM inventory WHERE product_id = %s"
        update_data_to_db(delete_inventory_sql, (product_id,))

        # 從 product_master 表中刪除
        delete_product_sql = "DELETE FROM product_master WHERE product_id = %s"
        update_data_to_db(delete_product_sql, (product_id,))

        return {
            "message": "產品刪除成功",
            "product_id": product_id
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] 產品刪除失敗: {e}")
        raise HTTPException(status_code=500, detail="產品刪除失敗")